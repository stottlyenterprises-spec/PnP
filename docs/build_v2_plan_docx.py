from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import re

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "PROGRESS_NOT_PERFECTION_VERSION_2_PLAN.md"
OUTPUT = ROOT / "Progress_Not_Perfection_Version_2_Plan.docx"

GREEN = RGBColor(25, 74, 57)
DARK = RGBColor(28, 47, 39)
MUTED = RGBColor(91, 109, 101)
PALE = "E8F0EB"
LINE = "CBD9D1"


def font(run, size=10.5, bold=False, italic=False, color=DARK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 7, GREEN),
        ("Heading 2", 13, 12, 5, GREEN),
        ("Heading 3", 11.5, 9, 4, DARK),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.font.color.rgb = DARK
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.1


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Progress, Not Perfection  |  Version 2   ")
    font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr_text, fld_char2))


def add_inline_markup(paragraph, text):
    pieces = re.split(r"(\*\*.*?\*\*)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            font(run, bold=True)
        else:
            run = paragraph.add_run(piece)
            font(run)


def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "260")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = ""
    hr = header.add_run("PNP  /  PRODUCT PLAN")
    font(hr, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Progress, Not Perfection")
    font(r, size=28, bold=True, color=GREEN, name="Aptos Display")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Version 2 Product and Build Plan")
    font(r, size=16, bold=True, color=DARK, name="Aptos Display")

    meta = doc.add_table(rows=2, cols=2)
    meta.autofit = False
    for row in meta.rows:
        row.cells[0].width = Inches(3.25)
        row.cells[1].width = Inches(3.25)
    values = [
        ("STATUS", "Product plan"),
        ("DATE", "July 26, 2026"),
        ("PRODUCT", "Progress, Not Perfection"),
        ("ASSISTANT", "D.E.E.D.S."),
    ]
    for cell, (label, value) in zip([c for row in meta.rows for c in row.cells], values):
        shade(cell, PALE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        a = p.add_run(label + "\n")
        font(a, size=7.5, bold=True, color=MUTED)
        b = p.add_run(value)
        font(b, size=10, bold=True, color=DARK)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(16)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.left_indent = Inches(0.18)
    lead.paragraph_format.right_indent = Inches(0.18)
    rr = lead.add_run(
        "Version 2 turns PNP from a capable personal dashboard into a dependable "
        "personal operating system for web, iPhone, and Android."
    )
    font(rr, size=12, bold=True, color=GREEN)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    started = False
    current_num_id = None
    previous_numbered = False
    for raw in lines:
        line = raw.strip()
        if line == "## Executive summary":
            started = True
            p = doc.add_paragraph("Executive summary", style="Heading 1")
            continue
        if not started or not line:
            previous_numbered = False
            continue
        if line.startswith("**Working status:**") or line.startswith("**Date:**") or line.startswith("**Product:**") or line.startswith("**Assistant:**"):
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            previous_numbered = False
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            previous_numbered = False
        elif re.match(r"^\d+\.\s", line):
            if not previous_numbered:
                current_num_id = new_numbering_id(doc)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            apply_numbering(p, current_num_id)
            add_inline_markup(p, re.sub(r"^\d+\.\s+", "", line))
            previous_numbered = True
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, line[2:])
            previous_numbered = False
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line[2:-2])
            font(r, bold=True, color=GREEN)
            previous_numbered = False
        else:
            p = doc.add_paragraph()
            p.paragraph_format.widow_control = True
            add_inline_markup(p, line)
            previous_numbered = False

    doc.core_properties.title = "Progress, Not Perfection: Version 2 Product and Build Plan"
    doc.core_properties.subject = "Version 2 product, mobile, data, and D.E.E.D.S. roadmap"
    doc.core_properties.author = "Progress, Not Perfection"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
